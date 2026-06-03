import markdown
from docx import Document
from docx.shared import Inches
import re

def markdown_to_docx(markdown_file, docx_file):
    # Read the markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()
    
    # Create a new Document
    doc = Document()
    
    # Process the markdown line by line
    lines = markdown_text.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Handle headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], level=5)
        elif line.startswith('###### '):
            doc.add_heading(line[7:], level=6)
        # Handle horizontal rules
        elif line.strip() == '---' or line.strip() == '***' or line.strip() == '___':
            doc.add_page_break()
        # Handle code blocks (simplified)
        elif line.startswith('```'):
            # Skip the opening ``` and collect code lines
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.add_run(code_text).font.name = 'Courier New'
                p.style = 'No Spacing'
        # Handle regular paragraphs
        else:
            if line.strip():  # Non-empty line
                p = doc.add_paragraph()
                # Handle bold and italic (simple version)
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        p.add_run(part[2:-2]).bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        p.add_run(part[1:-1]).italic = True
                    else:
                        if part:  # Don't add empty strings
                            p.add_run(part)
            else:  # Empty line
                doc.add_paragraph('')
        
        i += 1
    
    # Save the document
    doc.save(docx_file)
    print(f"Converted {markdown_file} to {docx_file}")

if __name__ == "__main__":
    markdown_to_docx("PROGRAMMING_EXPLANATION.md", "PROGRAMMING_EXPLANATION.docx")